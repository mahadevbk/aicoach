import streamlit as st
import cv2
import numpy as np
import os
import json
import mediapipe as mp
from mediapipe.tasks import python
from mediapipe.tasks.python import vision
import tempfile
import zipfile
import io
import urllib.request
import subprocess
import time
import generate_brief
import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from fpdf import FPDF
from fpdf.enums import XPos, YPos

# ============================================================================
# 1. CORE UTILITIES & ENCODERS
# ============================================================================

class NumpyEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, np.integer): return int(obj)
        if isinstance(obj, np.floating): return float(obj)
        if isinstance(obj, np.ndarray): return obj.tolist()
        if isinstance(obj, np.bool_): return bool(obj)
        return super(NumpyEncoder, self).default(obj)

THIN_JOINT_SPEC = (0, 255, 0)  # Bright Neon Green
THIN_BONE_SPEC = (255, 255, 255) # Pure White

def apply_advanced_slow_mo(input_path, output_path, impact_frame, fps):
    """
    Creates a smooth ramped slow-mo using segment-based processing.
    More robust than single-pass setpts for long or complex videos.
    """
    import subprocess
    
    # Calculate timestamps
    t_impact = impact_frame / fps
    t_start = max(0, t_impact - 0.5)
    t_end = t_impact + 0.5
    
    # Complex filter: split into 3 parts, slow down middle part, then concat
    # We use minterpolate ONLY on the middle part to save resources and ensure stability
    filter_complex = (
        f"[0:v]trim=start=0:end={t_start},setpts=PTS-STARTPTS,fps=60[v1]; "
        f"[0:v]trim=start={t_start}:end={t_end},setpts=5*(PTS-STARTPTS),minterpolate=fps=60:mi_mode=mci:mc_mode=obmc:me_mode=bidir[v2]; "
        f"[0:v]trim=start={t_end},setpts=PTS-STARTPTS,fps=60[v3]; "
        f"[v1][v2][v3]concat=n=3:v=1:a=0[outv]"
    )
    
    cmd = [
        'ffmpeg', '-y', '-i', input_path,
        '-filter_complex', filter_complex,
        '-map', '[outv]',
        '-c:v', 'libx264', '-pix_fmt', 'yuv420p', '-preset', 'ultrafast', '-crf', '20',
        '-movflags', '+faststart',
        output_path
    ]
    
    try:
        subprocess.run(cmd, check=True, capture_output=True, text=True)
    except subprocess.CalledProcessError as e:
        # Fallback to simple frame duplication if minterpolate fails
        fallback_filter = (
            f"[0:v]trim=start=0:end={t_start},setpts=PTS-STARTPTS,fps=60[v1]; "
            f"[0:v]trim=start={t_start}:end={t_end},setpts=5*(PTS-STARTPTS),fps=60[v2]; "
            f"[0:v]trim=start={t_end},setpts=PTS-STARTPTS,fps=60[v3]; "
            f"[v1][v2][v3]concat=n=3:v=1:a=0[outv]"
        )
        cmd_fallback = [
            'ffmpeg', '-y', '-i', input_path,
            '-filter_complex', fallback_filter,
            '-map', '[outv]',
            '-c:v', 'libx264', '-pix_fmt', 'yuv420p', '-preset', 'ultrafast',
            output_path
        ]
        subprocess.run(cmd_fallback, check=True)
        
    return output_path

# ============================================================================
# 2. REPORT GENERATION FUNCTIONS (Original)
# ============================================================================

def generate_sport_specific_prompt(sport, action, brief_content):
    """
    Generates a sport and action-specific coaching prompt dynamically.
    
    Args:
        sport: Sport name (e.g., "TENNIS", "GOLF", "GYM")
        action: Action type (e.g., "SERVE", "DRIVE", "SQUAT")
        brief_content: The telemetry brief text
    
    Returns:
        Tuple of (instructions, enhanced_brief)
    """
    
    # Sport-specific context library
    SPORT_COACHING_CONTEXT = {
        "TENNIS": {
            "discipline": "Racket Sport",
            "primary_chain": "Legs → Hips → Trunk → Shoulder → Elbow → Wrist",
            "key_metrics": ["racket_speed", "kinetic_chain_sequencing", "trunk_rotation", "hip_shoulder_separation"],
            "common_issues": ["poor hip-shoulder separation", "early arm movement", "insufficient trunk rotation", "foot placement"],
            "performance_indicators": ["first serve percentage", "racket head speed", "ball impact timing", "follow-through completion"],
            "injury_risks": ["rotator cuff strain", "lower back injury", "ankle sprain"],
        },
        "PADEL": {
            "discipline": "Racket Sport",
            "primary_chain": "Legs → Hips → Trunk → Shoulder → Elbow → Wrist",
            "key_metrics": ["racket_speed", "kinetic_chain_sequencing", "shoulder_stability", "wrist_control"],
            "common_issues": ["excessive wrist movement", "poor footwork", "uncontrolled arm swing", "weak loading position"],
            "performance_indicators": ["shot consistency", "racket speed", "body control", "positioning"],
            "injury_risks": ["tennis elbow", "shoulder impingement", "wrist strain"],
        },
        "PICKLEBALL": {
            "discipline": "Racket Sport",
            "primary_chain": "Legs → Hips → Trunk → Shoulder → Elbow → Wrist",
            "key_metrics": ["racket_control", "wrist_stability", "body_balance", "footwork"],
            "common_issues": ["excessive arm movement", "poor stance", "wrist instability", "timing issues"],
            "performance_indicators": ["shot accuracy", "control", "consistency", "positioning"],
            "injury_risks": ["wrist strain", "elbow strain", "shoulder strain"],
        },
        "BADMINTON": {
            "discipline": "Racket Sport",
            "primary_chain": "Legs → Hips → Trunk → Shoulder → Elbow → Wrist",
            "key_metrics": ["racket_speed", "shuttle_control", "footwork", "court_positioning"],
            "common_issues": ["poor footwork", "late timing", "weak loading", "incorrect grip transition"],
            "performance_indicators": ["shuttle speed", "accuracy", "footwork efficiency", "court coverage"],
            "injury_risks": ["ankle injury", "shoulder strain", "knee strain"],
        },
        "SQUASH": {
            "discipline": "Racket Sport",
            "primary_chain": "Legs → Hips → Trunk → Shoulder → Elbow → Wrist",
            "key_metrics": ["racket_speed", "court_positioning", "footwork", "shot_selection"],
            "common_issues": ["poor positioning", "late swing", "insufficient rotation", "weak footwork"],
            "performance_indicators": ["shot accuracy", "court control", "movement efficiency", "rally performance"],
            "injury_risks": ["shoulder strain", "knee injury", "ankle sprain"],
        },
        "GOLF": {
            "discipline": "Striking Sport",
            "primary_chain": "Legs → Hips → Trunk → Shoulders → Arms → Club",
            "key_metrics": ["club_head_speed", "hip_shoulder_separation", "swing_plane", "weight_transfer"],
            "common_issues": ["poor hip rotation", "early wrist release", "sway", "insufficient weight transfer"],
            "performance_indicators": ["ball_distance", "accuracy", "consistency", "launch_angle"],
            "injury_risks": ["lower back strain", "elbow strain", "shoulder strain"],
        },
        "GYM": {
            "discipline": "Strength/Conditioning",
            "primary_chain": "Ground → Legs → Core → Upper Body",
            "key_metrics": ["movement_symmetry", "range_of_motion", "stability", "alignment"],
            "common_issues": ["asymmetrical loading", "poor form", "insufficient depth", "core instability"],
            "performance_indicators": ["strength gains", "form quality", "stability", "symmetry"],
            "injury_risks": ["lower back strain", "knee strain", "shoulder impingement"],
        },
        "YOGA": {
            "discipline": "Flexibility/Mind-Body",
            "primary_chain": "Ground → Alignment → Breathing → Flow",
            "key_metrics": ["alignment", "stability", "flexibility", "balance"],
            "common_issues": ["poor alignment", "over-extension", "instability", "breath control"],
            "performance_indicators": ["form quality", "balance", "flexibility", "mind-body connection"],
            "injury_risks": ["overstretching", "wrist strain", "shoulder impingement"],
        },
        "SOCCER": {
            "discipline": "Kicking Sport",
            "primary_chain": "Standing Leg → Hips → Trunk → Kicking Leg → Foot",
            "key_metrics": ["ball_speed", "accuracy", "leg_power", "body_balance"],
            "common_issues": ["poor stance", "insufficient hip rotation", "weak follow-through", "balance issues"],
            "performance_indicators": ["ball_speed", "accuracy", "consistency", "placement"],
            "injury_risks": ["groin strain", "knee injury", "ankle injury"],
        },
        "BOXING/MMA": {
            "discipline": "Combat Sport",
            "primary_chain": "Legs → Hips → Core → Shoulders → Arms → Fists",
            "key_metrics": ["punch_speed", "hip_rotation", "balance", "footwork"],
            "common_issues": ["dropped hands", "poor footwork", "telegraphed movements", "power loss"],
            "performance_indicators": ["punch_speed", "combination_flow", "defense", "footwork_efficiency"],
            "injury_risks": ["hand injury", "rotator cuff strain", "cervical strain"],
        },
        "ATHLETICS/RUNNING": {
            "discipline": "Movement Science",
            "primary_chain": "Ground → Legs → Core → Upper Body → Arms",
            "key_metrics": ["stride_length", "cadence", "stability", "propulsion"],
            "common_issues": ["overstriding", "poor cadence", "instability", "asymmetry"],
            "performance_indicators": ["speed", "efficiency", "consistency", "endurance"],
            "injury_risks": ["knee injury", "shin splints", "hip strain"],
        },
        "BASEBALL": {
            "discipline": "Throwing Sport",
            "primary_chain": "Legs → Hips → Trunk → Shoulder → Elbow → Wrist",
            "key_metrics": ["throwing_velocity", "hip_shoulder_separation", "arm_angle", "follow_through"],
            "common_issues": ["early arm movement", "poor stride", "insufficient hip rotation", "arm lag"],
            "performance_indicators": ["throw_velocity", "accuracy", "consistency", "arm health"],
            "injury_risks": ["shoulder impingement", "ulnar collateral ligament strain", "rotator cuff strain"],
        },
        "AMERICAN FOOTBALL": {
            "discipline": "Throwing Sport",
            "primary_chain": "Legs → Hips → Trunk → Shoulder → Elbow → Wrist",
            "key_metrics": ["throwing_velocity", "accuracy", "release_point", "footwork"],
            "common_issues": ["poor stance", "timing issues", "weak release", "footwork errors"],
            "performance_indicators": ["pass_completion", "ball_velocity", "accuracy", "consistency"],
            "injury_risks": ["shoulder strain", "elbow strain", "knee injury"],
        },
        "ICE HOCKEY": {
            "discipline": "Racket Sport (Stick-based)",
            "primary_chain": "Legs → Hips → Trunk → Shoulders → Arms → Stick",
            "key_metrics": ["shot_speed", "accuracy", "body_balance", "footwork"],
            "common_issues": ["poor weight transfer", "weak loading", "timing issues", "stick control"],
            "performance_indicators": ["shot_speed", "accuracy", "consistency", "ice mobility"],
            "injury_risks": ["shoulder strain", "knee injury", "ankle injury"],
        },
        "TABLE TENNIS": {
            "discipline": "Racket Sport",
            "primary_chain": "Legs → Hips → Trunk → Shoulder → Elbow → Wrist",
            "key_metrics": ["paddle_speed", "footwork", "timing", "spin_control"],
            "common_issues": ["timing issues", "poor footwork", "weak positioning", "inconsistent contact"],
            "performance_indicators": ["shot_speed", "spin_production", "accuracy", "consistency"],
            "injury_risks": ["wrist strain", "elbow strain", "shoulder strain"],
        },
        "MARTIAL ARTS": {
            "discipline": "Combat Sport",
            "primary_chain": "Legs → Hips → Core → Shoulders → Arms → Limbs",
            "key_metrics": ["strike_speed", "balance", "hip_rotation", "footwork"],
            "common_issues": ["poor balance", "telegraphed strikes", "weak hip engagement", "footwork issues"],
            "performance_indicators": ["strike_speed", "accuracy", "balance", "technique_quality"],
            "injury_risks": ["hand injury", "rotator cuff strain", "knee injury"],
        },
    }
    
    # Get sport context, default if not in library
    context = SPORT_COACHING_CONTEXT.get(sport.upper(), {
        "discipline": "General Sport",
        "primary_chain": "Standard biomechanical chain",
        "key_metrics": ["general_performance_metrics"],
        "common_issues": ["performance_optimization"],
        "performance_indicators": ["overall_quality"],
        "injury_risks": ["general_injury_prevention"],
    })
    
    # Build dynamic prompt
    base_instructions = f"""
Act as both a professional {context['discipline'].lower()} coach and biomechanical engineer.

You are analyzing a {sport.upper()} {action.upper()} based on motion capture telemetry data. Be encouraging and positive

SPORT & ACTION CONTEXT:
- Sport: {sport.upper()}
- Action: {action.upper()}
- Primary kinetic chain: {context['primary_chain']}
- Key performance metrics for this action: {', '.join(context['key_metrics'])}
- Common technical issues in this action: {', '.join(context['common_issues'])}
- Performance indicators to assess: {', '.join(context['performance_indicators'])}
- Injury prevention focus: {', '.join(context['injury_risks'])}

MANDATORY REPORT STRUCTURE:
1. Executive Summary (2-3 sentences on overall {action.lower()} quality). Mention dominant side here.
2. Specialist coach's analysis, strengths and weaknesses.
3. Performance Scores table: 
   - Category | Score/100 | vs Benchmark | Assessment
   - Include scores for: Technique, Timing, Power, Balance/Stability, Follow-Through, Kinetic Chain, Consistency
4. Key Measurements table: 
   - Metric | Value | Benchmark | Status (Good/Attention/Issue)
   - Compare all provided measurements to benchmarks
4. Movement Progression Analysis:
   - Describe the movement from START → MID-POINT → FINISH
   - Identify timing of key events
   - Assess smoothness and efficiency
6. Kinetic Chain Assessment:
   - Analyze sequential energy transfer through: {context['primary_chain']}
   - Identify any breaks or delays in the chain
   - Assess efficiency of each segment
7. Critical Technical Issues (max 3):
   - Issue | Impact | Severity (HIGH/MEDIUM/LOW)
   - Explain the specific biomechanical problem
   - Note if it affects performance vs safety
8. Coaching Recommendations (max 5):
   - Recommendation | Priority (HIGH/MEDIUM/LOW) | Specific Cue | Progressive Drill
   - Focus on: {', '.join(context['common_issues'])}
9. Injury Risk Assessment:
   - Identify any movement patterns suggesting injury risk
   - Areas of concern: {', '.join(context['injury_risks'])}
   - Recommendations for prevention
10. Summary:
   - One sentence on current level
   - Next 3 priorities for improvement
   - Estimated timeline to see improvement

CRITICAL REQUIREMENTS:
- Always reference benchmark values from the brief when available
- Identify bilateral asymmetries (left vs right) and their impact
- Explain WHY issues matter, not just WHAT to do
- Distinguish between TIMING problems vs STRENGTH/MOBILITY problems
- Base all recommendations strictly on the telemetry data provided
- Use professional, direct language but remain encouraging
- Note any data quality limitations that affect interpretation

TONE: Professional, technical but accessible, positive but honest.
LENGTH: 900-1100 words
START WITH: "VECTOR VICTOR AI - BIO MECHANICAL ANALYSIS"
"""
    
    if action.upper() == "GENERAL RALLY":
        base_instructions += """
SPECIAL INSTRUCTION FOR GENERAL RALLY:
The telemetry data includes 'Detected Actions' with specific frame numbers in the metadata. 
1. Cross-reference the vector peaks at these frames with the action labels (e.g., Forehand vs Backhand).
2. Analyze the transition efficiency between these detected actions.
3. Compare the biomechanics of each individual shot detected within this single video.
"""

    return base_instructions

def generate_pro_report(brief_content, sport="GENERAL", action="MOVEMENT"):
    """
    Enhanced report generation using the 2026 flagship budget model: Gemini 3.1 Flash-Lite.
    
    Args:
        brief_content (str): The telemetry brief text from generate_brief()
        sport (str): Sport name (TENNIS, GOLF, GYM, etc.)
        action (str): Action type (SERVE, DRIVE, SQUAT, etc.)
    
    Returns:
        str: Generated coaching report text
    """
    try:
        # POINT TO THE NEW 2026 STANDARD
        # This replaces the old loop that looked for 1.5-pro
        selected_model = 'models/gemini-3.1-flash-lite-preview'
        
        # Initialize model
        report_model = genai.GenerativeModel(selected_model)
        
        # Generate sport-specific instructions (using your existing prompt helper)
        instructions = generate_sport_specific_prompt(sport, action, brief_content)
        
        # Use a consistent generation config for high-quality, long-form output
        generation_config = {
            "temperature": 0.7,
            "top_p": 0.95,
            "max_output_tokens": 4096,
        }
        
        # Call API with enhanced prompt and brief content
        response = report_model.generate_content(
            [instructions, brief_content],
            generation_config=generation_config
        )
        
        return response.text + """

Disclaimer                         
                                                                                                                    
This report is generated via automated computer-vision telemetry and is intended for performance-tracking and   
educational purposes only. Metrics (angles, velocities, and synchronization) are algorithmic estimates and do   
not constitute medical advice or a clinical diagnosis. Users should consult with a certified professional coach 
or healthcare provider before initiating new training intensities. Vector Victor AI assumes no liability for    
injury or performance outcomes resulting from the application of this data"""
    except Exception as e:
        # Check if the error is due to the model name (in case of regional rollout delays)
        return f"⚠️ AI Generation Error: {str(e)}"
 


def create_docx_report(text, sport_name, action, hand):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(8)
    
    # Main Title - Bitcount
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("VECTOR VICTOR AI")
    run.font.name = 'Bitcount Prop Single'
    run.bold = True
    run.font.size = Pt(10)
    
    # Subtitle - Arial 8pt Bold
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    s_run = subtitle.add_run("BIOMECHANICAL ANALYSIS REPORT")
    s_run.font.name = 'Arial'
    s_run.bold = True
    s_run.font.size = Pt(8)
    
    # Metadata - Arial 8pt
    p_date = doc.add_paragraph()
    p_date.paragraph_format.space_before = Pt(6)
    date_label = p_date.add_run("Date: ")
    date_label.font.name = 'Arial'
    date_label.font.size = Pt(8)
    date_label.bold = True
    date_value = p_date.add_run(time.strftime('%Y-%m-%d'))
    date_value.font.name = 'Arial'
    date_value.font.size = Pt(8)

    p_sport = doc.add_paragraph()
    sport_label = p_sport.add_run("Sport: ")
    sport_label.font.name = 'Arial'
    sport_label.font.size = Pt(8)
    sport_label.bold = True
    sport_value = p_sport.add_run(sport_name)
    sport_value.font.name = 'Arial'
    sport_value.font.size = Pt(8)
    
    p_activity = doc.add_paragraph()
    action_label = p_activity.add_run("Activity: ")
    action_label.font.name = 'Arial'
    action_label.font.size = Pt(8)
    action_label.bold = True
    action_value = p_activity.add_run(action)
    action_value.font.name = 'Arial'
    action_value.font.size = Pt(8)
    
    p_hand = doc.add_paragraph()
    hand_label = p_hand.add_run("Dominant Hand: ")
    hand_label.font.name = 'Arial'
    hand_label.font.size = Pt(8)
    hand_label.bold = True
    hand_value = p_hand.add_run(hand)
    hand_value.font.name = 'Arial'
    hand_value.font.size = Pt(8)
    
    # Add spacing after metadata
    doc.add_paragraph()
    
    # Access the footer of the first section (assuming a single section document)
    section = doc.sections[0]
    footer = section.footer
    
    # Add page number field and disclaimer text
    footer_paragraph = footer.add_paragraph()
    footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add disclaimer text
    run_disclaimer = footer_paragraph.add_run("Automated biomechanical analysis for educational use only. Consult a professional for clinical or medical assessment. | Page ")
    run_disclaimer.font.size = Pt(6)
    run_disclaimer.font.name = 'Arial'
    
    # Add PAGE field code using proper python-docx XML manipulation
    run_page = footer_paragraph.add_run()
    run_page.font.size = Pt(6)
    run_page.font.name = 'Arial'
    
    # Create field elements for PAGE
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    
    run_page._r.append(fldChar1)
    run_page._r.append(instrText)
    run_page._r.append(fldChar2)
    
    # Add " of " text
    run_of = footer_paragraph.add_run(" of ")
    run_of.font.size = Pt(6)
    run_of.font.name = 'Arial'
    
    # Add NUMPAGES field code
    run_numpages = footer_paragraph.add_run()
    run_numpages.font.size = Pt(6)
    run_numpages.font.name = 'Arial'
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'begin')
    
    instrText2 = OxmlElement('w:instrText')
    instrText2.set(qn('xml:space'), 'preserve')
    instrText2.text = 'NUMPAGES'
    
    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')
    
    run_numpages._r.append(fldChar3)
    run_numpages._r.append(instrText2)
    run_numpages._r.append(fldChar4)

    # Add the actual report content
    doc.add_paragraph()  # Add spacing after metadata
    
    # Parse and add the report text content
    lines = text.split('\n')
    current_table = None
    current_table_cols = 0
    
    for line in lines:
        stripped = line.strip()
        if stripped:  # Only add non-empty lines
            # Check if it's a heading (lines with #)
            if stripped.startswith('#'):
                current_table = None  # Reset table context
                # Remove # symbols and add as heading
                heading_text = stripped.lstrip('#').strip()
                level = len(stripped) - len(stripped.lstrip('#'))
                # Limit heading levels to valid range (1-2)
                level = min(level, 2)
                if level > 0:
                    heading = doc.add_heading(heading_text, level=level)
                    # Ensure heading uses Arial 8pt bold
                    for run in heading.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(8)
                        run.bold = True
                else:
                    p = doc.add_paragraph(heading_text, style='Normal')
                    for run in p.runs:
                        run.font.name = 'Arial'
                        run.font.size = Pt(8)
                        run.bold = True
            # Check if it's a table row (contains |)
            elif '|' in stripped and '---' not in stripped:
                # Table handling - extract cells
                cells = [c.strip() for c in stripped.split('|') if c.strip()]
                if cells:
                    # Check if we need a new table or continue existing
                    if current_table is None or len(cells) != current_table_cols:
                        current_table = doc.add_table(rows=1, cols=len(cells))
                        current_table.style = 'Light Grid Accent 1'
                        current_table_cols = len(cells)
                    else:
                        # Add row to existing table
                        current_table.add_row()
                    
                    # Fill cells with text and proper formatting
                    row_idx = len(current_table.rows) - 1
                    for i, cell_text in enumerate(cells):
                        cell = current_table.rows[row_idx].cells[i]
                        cell.text = cell_text
                        # Format cell text
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.name = 'Arial'
                                run.font.size = Pt(8)
            # Regular paragraph
            else:
                current_table = None  # Reset table context
                p = doc.add_paragraph(stripped, style='Normal')
                p.paragraph_format.space_after = Pt(6)
                # Ensure paragraph uses Arial 8pt
                for run in p.runs:
                    run.font.name = 'Arial'
                    run.font.size = Pt(8)
        else:
            current_table = None  # Reset table context on empty line
            # Empty line = paragraph break
            doc.add_paragraph()

    # Save document to BytesIO
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

class PDFReport(FPDF):
    def __init__(self):
        super().__init__(format='A4')
        self.custom_font_active = False
        try:
            # Check local folder first (e.g. if uploaded to GitHub)
            local_font = "BitcountPropSingle-VariableFont_CRSV,ELSH,ELXP,slnt,wght.ttf"
            if os.path.exists(local_font):
                font_path = local_font
            else:
                font_path = os.path.join(tempfile.gettempdir(), "BitcountPropSingle-VF.ttf")
                if not os.path.exists(font_path):
                    # Fallback to a similar font if local not found and download fails
                    urllib.request.urlretrieve("https://github.com/google/fonts/raw/main/ofl/robotoflex/RobotoFlex%5BGRAD%2CXOPQ%2CXTRA%2CYOPQ%2CYTLC%2CYTAS%2CYTDE%2CYTFI%2Copsz%2Cslnt%2Cwdth%2Cwght%5D.ttf", font_path)
            self.add_font("CustomFont", "", font_path)
            self.add_font("CustomFont", "B", font_path) # Register same file for bold to prevent crashes
            self.custom_font_active = True
        except:
            self.custom_font_active = False

    def header(self):
        f_name = "CustomFont" if self.custom_font_active else "helvetica"
        self.set_font(f_name, 'B', 10) # Bold and size 10 for main title
        self.set_text_color(0, 180, 255)
        self.cell(0, 10, 'VECTOR VICTOR AI', border=False, align='C', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        self.set_font(f_name, 'B', 8) # Bold for subtitle
        self.set_text_color(100, 100, 100)
        self.cell(0, 5, 'BIO MECHANICAL ANALYSIS REPORT', border=False, align='C', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        self.ln(2)
        self.set_draw_color(0, 180, 255)
        self.line(10, 28, 200, 28)
        self.ln(5)

    def footer(self):
        self.set_y(-15)
        self.set_font('helvetica', '', 6) # Size 6 text
        self.set_text_color(128, 128, 128)
        # Add disclaimer and page number
        disclaimer = "Automated biomechanical analysis for educational use only. Consult a professional for clinical or medical assessment."
        self.cell(0, 10, f"{disclaimer} | Page {self.page_no()}/{{nb}}", align='C')

def create_pdf_report(text, sport_name, action, hand):
    def clean_for_pdf(s):
        try:
            return s.encode('latin-1', 'replace').decode('latin-1').replace('?', '')
        except:
            return "".join([c for c in s if ord(c) < 128])

    pdf = PDFReport()
    pdf.alias_nb_pages()
    pdf.add_page()
    # Body font is Helvetica (universally available sans-serif)
    f_main = "helvetica"
    
    pdf.set_font(f_main, 'B', 8)
    pdf.write(6, "Date: ")
    pdf.set_font(f_main, '', 8)
    pdf.write(6, f"{time.strftime('%Y-%m-%d %H:%M:%S')}\n")

    pdf.set_font(f_main, 'B', 8)
    pdf.write(6, "Sport: ")
    pdf.set_font(f_main, '', 8)
    pdf.write(6, f"{clean_for_pdf(sport_name)}\n")
    
    pdf.set_font(f_main, 'B', 8)
    pdf.write(6, "Activity: ")
    pdf.set_font(f_main, '', 8)
    pdf.write(6, f"{clean_for_pdf(action)}\n")
    
    pdf.set_font(f_main, 'B', 8)
    pdf.write(6, "Hand dominance: ")
    pdf.set_font(f_main, '', 8)
    pdf.write(6, f"{clean_for_pdf(hand)}\n")
    pdf.ln(4)

    lines = text.split('\n')
    table_data = []
    in_table = False

    for line in lines:
        clean_line = clean_for_pdf(line.strip())
        if '|' in clean_line:
            if '---' in clean_line: continue
            cells = [c.strip() for c in clean_line.split('|') if c.strip()]
            if cells:
                table_data.append(cells)
                in_table = True
            continue
        else:
            if in_table and table_data:
                # Use fpdf2 table feature for automatic wrapping
                with pdf.table(borders_layout="ALL", gutter_height=1, text_align="LEFT") as table:
                    for r_idx, row_data in enumerate(table_data):
                        row = table.row()
                        if r_idx == 0:
                            pdf.set_font(f_main, 'B', 8)
                        else:
                            pdf.set_font(f_main, '', 8)
                        for val in row_data:
                            row.cell(val)
                pdf.ln(4)
                table_data = []
                in_table = False

        if not clean_line:
            pdf.ln(1)
            continue

        pdf.set_x(10)
        pdf.set_font(f_main, '', 8)
        if clean_line.startswith('### '):
            pdf.set_text_color(80, 80, 80)
            pdf.multi_cell(0, 6, clean_line.replace('### ', ''), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        elif clean_line.startswith('## '):
            pdf.set_font(f_main, 'B', 8) # Explicitly bold
            pdf.set_text_color(0, 0, 0)
            pdf.multi_cell(0, 6, clean_line.replace('## ', ''), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        elif clean_line.startswith('* ') or clean_line.startswith('- '):
            pdf.set_text_color(0, 0, 0)
            content = clean_line[2:]
            pdf.write(5, "  - ")
            if ':' in content:
                title, desc = content.split(':', 1)
                pdf.set_font(f_main, 'B', 8)
                pdf.write(5, title + ":")
                pdf.set_font(f_main, '', 8)
                pdf.write(5, desc + "\n")
            else:
                pdf.set_font(f_main, '', 8)
                pdf.write(5, content + "\n")
        else:
            pdf.set_font(f_main, '', 8)
            pdf.set_text_color(30, 30, 30)
            if '**' in clean_line:
                parts = clean_line.split('**')
                for idx_p, part in enumerate(parts):
                    if idx_p % 2 == 1: pdf.set_font(f_main, 'B', 8)
                    else: pdf.set_font(f_main, '', 8)
                    pdf.write(5, part)
                pdf.ln(5)
            else:
                pdf.multi_cell(0, 5, clean_line, align='L')
        pdf.ln(1)

    return bytes(pdf.output())

# ============================================================================
# 3. PAGE CONFIG & MOBILE CSS
# ============================================================================

st.set_page_config(
    page_title="Vector Victor AI Skeletonkey", 
    page_icon="🎾", 
    layout="wide", 
    initial_sidebar_state="collapsed"
)

# Setup Gemini
if "GEMINI_API_KEY" in st.secrets:
    genai.configure(api_key=st.secrets["GEMINI_API_KEY"])

import plotly.graph_objects as go
import plotly.express as px

st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Bitcount+Prop+Single&display=swap');
    :root {
        --neon-green: #ccff00;
        --glass-bg: rgba(255, 255, 255, 0.05);
        --glass-border: rgba(255, 255, 255, 0.1);
    }
    
    /* Global Font & Uppercase Override */
    .stApp, .stApp * { 
        font-family: 'Bitcount Prop Single', sans-serif !important; 
        font-weight: normal !important;
    }
    
    /* Global Uppercase for Interactive Elements */
    button, label, input, select, textarea, .stTabs [data-baseweb="tab"] p, .hero-sub, h4, [data-testid="stWidgetLabel"] p { 
        text-transform: uppercase !important;
        font-weight: normal !important;
    }

    /* Standardise File Uploader - Keep Font, Fix Overlap by disabling uppercase */
    [data-testid="stFileUploader"] * {
        text-transform: none !important;
        font-weight: normal !important;
    }
    
    .stApp { background: radial-gradient(circle at top right, #0f172a, #020617); color: #f8fafc; }
    .main { padding: 1rem 0.5rem; }
    
    /* Button Styling */
    button { 
        min-height: 48px !important;
        font-size: 14px !important;
        font-weight: normal !important;
    }
    .stButton > button { width: 100% !important; }

    /* Tabs Styling */
    .stTabs [data-baseweb="tab-list"] { gap: 0.5rem; display: flex; flex-wrap: wrap; justify-content: center; }
    .stTabs [data-baseweb="tab"] {
        flex: 1; min-width: 80px; text-align: center; height: 50px; 
        background: rgba(255, 255, 255, 0.05) !important; border-radius: 10px !important; 
        border: 1px solid rgba(255,255,255,0.1) !important;
    }
    .stTabs [aria-selected="true"] { background: rgba(204, 255, 0, 0.1) !important; border: 1px solid var(--neon-green) !important; }
    .stTabs [aria-selected="true"] p { color: var(--neon-green) !important; font-weight: normal !important; }

    /* Slider & Column Styling */
    .stSlider { padding: 1rem 0; }
    [data-testid="column"] { padding: 0 0.25rem; }
    div[data-testid="stSlider"] label p { color: var(--neon-green) !important; font-size: 1rem !important; letter-spacing: 1px; font-weight: normal !important; }
    div[data-testid="stThumbValue"] { color: #000 !important; background-color: var(--neon-green) !important; font-weight: normal !important; }
    div[data-baseweb="slider"] > div { background: var(--neon-green) !important; }
    
    /* Card Styling */
    .glass-card { background: var(--glass-bg); backdrop-filter: blur(12px); border: 1px solid var(--glass-border); border-radius: 24px; padding: 1.5rem; margin-bottom: 2rem; }
    .bento-card { background: rgba(255, 255, 255, 0.03); border: 1px solid var(--glass-border); border-radius: 20px; padding: 20px; text-align: center; margin-bottom: 10px; }

    /* Heading Styling */
    h1 { font-size: clamp(2rem, 8vw, 4rem) !important; font-weight: normal !important; background: linear-gradient(to right, #00f2fe, #4facfe); -webkit-background-clip: text; -webkit-text-fill-color: transparent; text-align: center; margin-bottom: 0px !important; text-transform: uppercase !important; }
    h4 { font-size: 0.85rem !important; color: var(--neon-green); margin-top: 2.5rem !important; margin-bottom: 0.5rem !important; letter-spacing: 2px; font-weight: normal !important; }
    .hero-sub { text-align: center; color: #94a3b8; font-size: 0.7rem; letter-spacing: 4px; margin-bottom: 2rem; font-weight: normal !important; }

    /* Specific Button Highlighting */
    div.stButton > button:has(div:contains("GENERATE AI COACHING REPORT")) { background: linear-gradient(135deg, #ffd700 0%, #daa520 50%, #b8860b 100%) !important; color: #000 !important; font-weight: normal !important; }
    </style>
    """, unsafe_allow_html=True)

# ============================================================================
# 4. BIOMECHANIC CALCULATORS (Original)
# ============================================================================

def calculate_3d_angle(p1, p2, p3):
    a, b, c = np.array([p1['x'], p1['y'], p1['z']]), np.array([p2['x'], p2['y'], p2['z']]), np.array([p3['x'], p3['y'], p3['z']])
    ba, bc = a - b, c - b
    cosine_angle = np.dot(ba, bc) / (np.linalg.norm(ba) * np.linalg.norm(bc) + 1e-6)
    return round(float(np.degrees(np.arccos(np.clip(cosine_angle, -1.0, 1.0)))), 1)

def get_dist(p1, p2):
    return np.linalg.norm(np.array([p1['x'], p1['y'], p1['z']]) - np.array([p2['x'], p2['y'], p2['z']]))

def get_midpoint(p1, p2):
    return {'x': (p1['x']+p2['x'])/2, 'y': (p1['y']+p2['y'])/2, 'z': (p1['z']+p2['z'])/2}

def calculate_lean(p1, p2, plane='sagittal'):
    v = np.array([p2['x']-p1['x'], p2['y']-p1['y'], p2['z']-p1['z']])
    if plane == 'sagittal': # Y-Z plane
        angle = np.degrees(np.arctan2(v[2], v[1]))
    else: # Coronal: X-Y plane
        angle = np.degrees(np.arctan2(v[0], v[1]))
    return round(float(angle), 1)

def interpolate_landmarks(raw_frames):
    total = len(raw_frames)
    if total == 0: return []
    frames = [f.copy() if f else None for f in raw_frames]
    for j in range(33):
        missing = [i for i in range(total) if not frames[i] or frames[i][j] is None]
        if not missing or len(missing) == total: continue
        for m_idx in missing:
            prev_idx = next((i for i in range(m_idx-1, -1, -1) if frames[i] and frames[i][j]), None)
            next_idx = next((i for i in range(m_idx+1, total) if frames[i] and frames[i][j]), None)
            if prev_idx is not None and next_idx is not None:
                p, n = frames[prev_idx][j], frames[next_idx][j]
                t = (m_idx - prev_idx) / (next_idx - prev_idx)
                val = {'x': p['x'] + t*(n['x']-p['x']), 'y': p['y'] + t*(n['y']-p['y']), 'z': p['z'] + t*(n['z']-p['z'])}
            elif prev_idx is not None: val = frames[prev_idx][j]
            elif next_idx is not None: val = frames[next_idx][j]
            else: val = {'x': 0.0, 'y': 0.0, 'z': 0.0}
            if not frames[m_idx]: frames[m_idx] = [{'x':0.0, 'y':0.0, 'z':0.0} for _ in range(33)]
            frames[m_idx][j] = val
    return frames

def detect_handedness(raw_frames):
    r_wrist_speeds, l_wrist_speeds = [], []
    for f in raw_frames:
        if not f: continue
        # Simple velocity approximation for wrist markers (16=right, 15=left)
        r_wrist_speeds.append(f[16]['x']) # Using horizontal position for simplicity
        l_wrist_speeds.append(f[15]['x'])
    
    r_var = np.var(r_wrist_speeds) if len(r_wrist_speeds) > 1 else 0
    l_var = np.var(l_wrist_speeds) if len(l_wrist_speeds) > 1 else 0
    return "right" if r_var > l_var else "left"

# Add this to Section 4: BIOMECHANIC CALCULATORS
ACTION_SIGNATURES = {
    "TENNIS": {
        "SERVE": {"metric": "r_wrist_speed", "threshold": 0.85, "min_elbow": 150, "min_shoulder": 110},
        "FOREHAND DRIVE": {"metric": "r_wrist_speed", "threshold": 0.5, "min_elbow": 80, "max_shoulder": 100},
        "BACKHAND DRIVE": {"metric": "l_wrist_speed", "threshold": 0.5, "min_elbow": 80},
        "FOREHAND SLICE": {"metric": "r_wrist_speed", "threshold": 0.4, "max_elbow": 110},
    }
}

def auto_detect_actions(metrics, sport, fps=30):
    """Detects frame indices where key actions occur based on telemetry peaks with priority and validation."""
    detected = []
    sport = "".join([c for c in sport if ord(c) < 128]).strip().upper()
    if sport not in ACTION_SIGNATURES: return detected
    
    # Get total frames from any available metric
    first_metric = next(iter(metrics.values()), [])
    total_frames = len(first_metric)
    if total_frames < 3: return detected

    # Suppression window: 8 seconds
    suppression_window = int(fps * 8)

    i = 1
    while i < total_frames - 1:
        candidates = []
        for action_name, sig in ACTION_SIGNATURES[sport].items():
            speeds = metrics.get(sig['metric'], [])
            if i >= len(speeds): continue
            
            val = speeds[i]
            if val is None or val <= sig['threshold']: continue
            
            # Local peak check (robust against None)
            prev_val = speeds[i-1] if (i > 0 and speeds[i-1] is not None) else 0
            next_val = speeds[i+1] if (i < total_frames-1 and speeds[i+1] is not None) else 0
            if not (val > prev_val and val > next_val): continue
            
            # Secondary verification
            passed_secondary = True
            # Determine which side to check based on the wrist metric side
            side_prefix = "r" if "r_wrist" in sig['metric'] else "l"
            
            # Elbow Check
            elbow_angles = metrics.get(f"{side_prefix}_elbow", [])
            if 'min_elbow' in sig:
                if i < len(elbow_angles) and elbow_angles[i] is not None:
                    if elbow_angles[i] < sig['min_elbow']: passed_secondary = False
                else: passed_secondary = False
            if 'max_elbow' in sig:
                if i < len(elbow_angles) and elbow_angles[i] is not None:
                    if elbow_angles[i] > sig['max_elbow']: passed_secondary = False
                else: passed_secondary = False
            
            # Shoulder Check (Abduction)
            shoulder_angles = metrics.get(f"{side_prefix}_shoulder_abduction", [])
            if 'min_shoulder' in sig:
                if i < len(shoulder_angles) and shoulder_angles[i] is not None:
                    if shoulder_angles[i] < sig['min_shoulder']: passed_secondary = False
                else: passed_secondary = False
            if 'max_shoulder' in sig:
                if i < len(shoulder_angles) and shoulder_angles[i] is not None:
                    if shoulder_angles[i] > sig['max_shoulder']: passed_secondary = False
                else: passed_secondary = False
                
            if passed_secondary:
                candidates.append({"action": action_name, "frame": i, "priority": sig['threshold']})
        
        if candidates:
            # Select action with the highest threshold (priority)
            best = max(candidates, key=lambda x: x['priority'])
            detected.append({"action": best['action'], "frame": int(best['frame'])})
            i += suppression_window # Temporal suppression: 8 seconds skip
        else:
            i += 1
            
    return detected

def build_pro_telemetry(raw_frames, sport_raw, action, event_frame, fps, camera_mode, handedness_override=None):
    total_frames = len(raw_frames)
    sport_clean = "".join([c for c in sport_raw if ord(c) < 128]).strip().upper()
    event_frame = max(0, min(event_frame, total_frames - 1))
    
    detected_handedness = detect_handedness(raw_frames)
    final_handedness = handedness_override if handedness_override else detected_handedness
    
    metrics = {
        "r_elbow": [], "l_elbow": [], "r_knee": [], "l_knee": [], "r_hip": [], "l_hip": [],
        "r_shoulder_abduction": [], "l_shoulder_abduction": [], "r_ankle": [], "l_ankle": [],
        "r_wrist_speed": [], "l_wrist_speed": [], "r_ankle_speed": [], "l_ankle_speed": [],
        "shoulder_z_diff": [], "hip_z_diff": [], "trunk_forward_lean": [], "trunk_lateral_lean": []
    }
    
    validation_warnings = []
    filtered_pos = []
    for i in range(total_frames):
        f_filt = {}
        for pt_idx in [15, 16, 27, 28]:
            win_x, win_y, win_z = [], [], []
            for d in [-1, 0, 1]:
                idx = max(0, min(total_frames - 1, i + d))
                if raw_frames[idx]:
                    win_x.append(raw_frames[idx][pt_idx]['x'])
                    win_y.append(raw_frames[idx][pt_idx]['y'])
                    win_z.append(raw_frames[idx][pt_idx]['z'])
            if win_x:
                f_filt[pt_idx] = {'x': float(np.median(win_x)), 'y': float(np.median(win_y)), 'z': float(np.median(win_z))}
            else:
                f_filt[pt_idx] = {'x': 0.0, 'y': 0.0, 'z': 0.0}
        filtered_pos.append(f_filt)

    last_valid_scale = 0.5
    prev_pts = [None] * 4 # RW, LW, RA, LA
    
    for idx_f, f in enumerate(raw_frames):
        if not f:
            for k in metrics: metrics[k].append(None)
            continue
            
        mid_s, mid_h = get_midpoint(f[11], f[12]), get_midpoint(f[23], f[24])
        current_scale = get_dist(mid_s, mid_h)
        scale = current_scale if current_scale >= 0.05 else last_valid_scale
        last_valid_scale = scale
            
        metrics["r_elbow"].append(calculate_3d_angle(f[12], f[14], f[16]))
        metrics["l_elbow"].append(calculate_3d_angle(f[11], f[13], f[15]))
        metrics["r_knee"].append(calculate_3d_angle(f[24], f[26], f[28]))
        metrics["l_knee"].append(calculate_3d_angle(f[23], f[25], f[27]))
        metrics["r_hip"].append(calculate_3d_angle(f[12], f[24], f[26]))
        metrics["l_hip"].append(calculate_3d_angle(f[11], f[23], f[25]))
        metrics["r_shoulder_abduction"].append(calculate_3d_angle(f[14], f[12], f[24]))
        metrics["l_shoulder_abduction"].append(calculate_3d_angle(f[13], f[11], f[23]))
        metrics["r_ankle"].append(calculate_3d_angle(f[26], f[28], f[30]))
        metrics["l_ankle"].append(calculate_3d_angle(f[25], f[27], f[29]))
        
        speed_keys = ["r_wrist_speed", "l_wrist_speed", "r_ankle_speed", "l_ankle_speed"]
        for i, pt_idx in enumerate([16, 15, 28, 27]):
            curr_pt = filtered_pos[idx_f][pt_idx]
            if prev_pts[i]:
                metrics[speed_keys[i]].append(round(get_dist(curr_pt, prev_pts[i]) / scale, 4))
            else:
                metrics[speed_keys[i]].append(0.0)
            prev_pts[i] = curr_pt
        
        metrics["shoulder_z_diff"].append(round(f[12]['z'] - f[11]['z'], 3))
        metrics["hip_z_diff"].append(round(f[24]['z'] - f[23]['z'], 3))
        metrics["trunk_forward_lean"].append(calculate_lean(mid_s, mid_h, 'sagittal'))
        metrics["trunk_lateral_lean"].append(calculate_lean(mid_s, mid_h, 'coronal'))

    for key in ["r_wrist_speed", "l_wrist_speed", "r_ankle_speed", "l_ankle_speed"]:
        if key in metrics:
            for i in range(len(metrics[key])):
                if metrics[key][i] is not None and metrics[key][i] > 1.5:
                    metrics[key][i] = round(((metrics[key][max(0, i-1)] or 0.0) + (metrics[key][min(total_frames-1, i+1)] or 0.0)) / 2, 4)

    racket_sports = ["TENNIS", "PADEL", "PICKLEBALL", "BADMINTON", "SQUASH"]
    if sport_clean in racket_sports:
        for k in ["r_ankle", "l_ankle", "r_ankle_speed", "l_ankle_speed"]: 
            if k in metrics: del metrics[k]

    def get_snapshot(idx):
        idx = max(0, min(idx, total_frames - 1))
        f = raw_frames[idx]
        if not f: return {}
        mid_s, mid_h = get_midpoint(f[11], f[12]), get_midpoint(f[23], f[24])
        hip_dist = get_dist(f[23], f[24])
        tilt = np.degrees(np.arctan2(f[12]['y'] - f[11]['y'], f[12]['x'] - f[11]['x']))
        if abs(tilt) > 90: tilt = tilt - 180 if tilt > 0 else tilt + 180
        sw_ratio = round(get_dist(f[27], f[28]) / (hip_dist + 1e-6), 4)

        return {
            "r_elbow_angle": calculate_3d_angle(f[12], f[14], f[16]), "l_elbow_angle": calculate_3d_angle(f[11], f[13], f[15]),
            "r_knee_angle": calculate_3d_angle(f[24], f[26], f[28]), "l_knee_angle": calculate_3d_angle(f[23], f[25], f[27]),
            "r_hip_angle": calculate_3d_angle(f[12], f[24], f[26]), "l_hip_angle": calculate_3d_angle(f[11], f[23], f[25]),
            "r_shoulder_abduction_angle": calculate_3d_angle(f[14], f[12], f[24]), "l_shoulder_abduction_angle": calculate_3d_angle(f[13], f[11], f[23]),
            "shoulder_tilt_deg": round(tilt, 1),
            "trunk_forward_lean": calculate_lean(mid_s, mid_h, 'sagittal'), "trunk_lateral_lean": calculate_lean(mid_s, mid_h, 'coronal'),
            "shoulder_z_diff": round(f[12]['z'] - f[11]['z'], 3), "hip_z_diff": round(f[24]['z'] - f[23]['z'], 3),
            "hip_shoulder_separation": round((f[12]['z'] - f[11]['z']) - (f[24]['z'] - f[23]['z']), 3),
            "r_wrist_above_r_shoulder": f[16]['y'] < f[12]['y'], "l_wrist_above_l_shoulder": f[15]['y'] < f[11]['y'],
            "feet_grounded": f[28]['y'] > 0.80 and f[27]['y'] > 0.80, "stance_width_ratio": sw_ratio if sw_ratio < 2.5 else None
        }

    output = {
        "sport": sport_clean, "action": action, "camera": camera_mode,
        "metadata": {
            "fps": fps, "total_frames": total_frames,
            "dominant_side": final_handedness,
            "coordinate_system": {"y_axis": "increases_downward", "z_axis": "depth_into_camera", "normalisation": "mediapipe_image_fraction_0_to_1"},
            "validation_warnings": validation_warnings
        },
        "metrics": {k: [v for v in metrics[k] if v is not None] for k in metrics if any(v is not None for v in metrics[k])},
        "event_snapshot": get_snapshot(event_frame), "phase_snapshots": {}, "speed_analysis": {}, "rotation_analysis": {}, "balance_stability": {}
    }

    phases = []
    if sport_clean in racket_sports: phases = [("trophy", -40), ("swing_start", -15), ("follow_through", 20)]
    elif sport_clean == "GOLF": phases = [("address", -80), ("top", -30), ("downswing", -12), ("follow", 25)]
    elif sport_clean == "GYM": phases = [("start", -45), ("midpoint", -22), ("finish", 30)]
    elif sport_clean == "YOGA": phases = [("approach", -30), ("exit", 30)]
    elif sport_clean == "SOCCER": phases = [("approach", -15), ("strike", 0), ("follow_through", 10)]
    elif sport_clean == "BOXING/MMA": phases = [("load", -10), ("impact", 0), ("recoil", 8)]
    elif sport_clean == "ATHLETICS/RUNNING": phases = [("drive", -5), ("extension", 0), ("recovery", 12)]
    elif sport_clean == "BASEBALL": phases = [("Leg Lift", -40), ("Stride Foot Plant", -15), ("Release/Impact", 0), ("Follow Through", +20)]
    elif sport_clean == "AMERICAN FOOTBALL": phases = [("Set/Drop", -20), ("Release/Kick", 0), ("Follow Through", +15)]
    elif sport_clean == "ICE HOCKEY": phases = [("Backswing", -12), ("Puck Impact", 0), ("Follow Through", +10)]
    elif sport_clean == "TABLE TENNIS": phases = [("Backswing", -8), ("Impact", 0), ("Recovery", +5)]
    elif sport_clean == "MARTIAL ARTS": phases = [("Chamber/Load", -10), ("Impact", 0), ("Reset", +8)]
    for name, p_off in phases: output["phase_snapshots"][name] = get_snapshot(event_frame + p_off)

    def analyze_speed(speed_series):
        clean_s = [s if s is not None else 0 for s in speed_series]
        peak_idx = np.argmax(clean_s)
        return {"peak_speed": round(float(np.max(clean_s)), 4), "peak_frame_offset": int(peak_idx - event_frame), "speed_at_event": round(clean_s[event_frame], 4)}

    output["speed_analysis"]["r_wrist"] = analyze_speed(metrics["r_wrist_speed"])
    output["speed_analysis"]["l_wrist"] = analyze_speed(metrics["l_wrist_speed"])

    def find_velocity_peak(series):
        vel = [abs(series[i] - series[i-1]) for i in range(1, len(series))]
        vel = [0] + vel
        start, end = max(0, event_frame-60), min(total_frames, event_frame+6)
        window = vel[start:end]
        return start + np.argmax(window) if window else event_frame

    shoulder_p = find_velocity_peak(metrics["shoulder_z_diff"])
    hip_p = find_velocity_peak(metrics["hip_z_diff"])
    output["rotation_analysis"] = {
        "hip_leads_shoulder": hip_p < shoulder_p, "hip_peak_offset": int(hip_p - event_frame), "shoulder_peak_offset": int(shoulder_p - event_frame),
        "x_factor_at_event": round(metrics["shoulder_z_diff"][event_frame] - metrics["hip_z_diff"][event_frame], 3)
    }
    
    win = [f[0]['x'] for f in raw_frames[max(0, event_frame-5):min(total_frames, event_frame+5)] if f]
    output["balance_stability"] = {"nose_x_variance": round(float(np.var(win)), 4) if win else 0}
    return output

# ============================================================================
# 5. UI HELPERS (Mobile Optimized)
# ============================================================================

def draw_modern_metric(label, value, delta, icon="⚡"):
    st.markdown(f"""
        <div class="bento-card">
            <div style="font-size: 0.8rem; color: #94a3b8; text-transform: uppercase; letter-spacing: 1px;">{label}</div>
            <div style="font-size: 1.8rem; font-weight: 900; color: #ccff00; margin: 10px 0;">{icon} {value}</div>
            <div style="font-size: 0.9rem; color: #38bdf8;">{delta}</div>
        </div>
    """, unsafe_allow_html=True)

def display_file_info(uploaded_file):
    if uploaded_file:
        file_size_mb = uploaded_file.size / (1024 * 1024)
        st.caption(f"📊 {uploaded_file.name} • {file_size_mb:.1f} MB")

def draw_mobile_metric_grid(metrics_dict):
    items = list(metrics_dict.items())
    for i in range(0, len(items), 2):
        col1, col2 = st.columns(2)
        with col1:
            label, (value, delta, icon) = items[i]
            draw_modern_metric(label, value, delta, icon)
        if i + 1 < len(items):
            with col2:
                label, (value, delta, icon) = items[i + 1]
                draw_modern_metric(label, value, delta, icon)

def get_sport_metrics(sport, metrics, sl1, kpis, s):
    # Dynamic calculations for Max Velocity
    speed_key = "wrist_speed"
    if any(x in sport for x in ["SOCCER", "ATHLETICS", "FOOTBALL"]):
        speed_key = "ankle_speed"
    max_v = max([v for v in metrics[speed_key] if v is not None] or [0])
    v_label = "Elite" if max_v > 20 else ("Optimal" if max_v > 10 else "Developing")
    
    # Efficiency based on trunk lateral lean
    lateral_lean_at_impact = metrics["trunk_lateral_lean"][sl1] if "trunk_lateral_lean" in metrics and sl1 < len(metrics["trunk_lateral_lean"]) else 0
    efficiency = max(0, min(100, 100 - (abs(lateral_lean_at_impact) * 3)))
    eff_label = "Elite" if efficiency > 85 else "Optimal"
    
    m_dict = {
        "Max Velocity": (f"{max_v:.1f}m/s", v_label, "⚡"),
        "Bio-Efficiency": (f"{efficiency:.0f}%", eff_label, "🧬")
    }
    
    if "GYM" in sport:
        depth = kpis.get('depth_ratio', 0)
        m_dict["Squat Depth"] = (f"{depth:.2f}", "Elite" if depth < 0.8 else "Developing", "📏")
    elif "GOLF" in sport:
        xf = kpis.get('max_x_factor', 0)
        m_dict["X-Factor"] = (f"{xf:.1f}°", "Elite" if xf > 40 else "Developing", "🔄")
    elif "YOGA" in sport:
        stab = kpis.get('stability', 0)*100
        m_dict["Stability"] = (f"{stab:.1f}%", "Elite" if stab > 95 else "Optimal", "🧘")
    elif any(x in sport for x in ["SOCCER", "BASEBALL"]):
        linkage = kpis.get("linkage_score", 0)
        m_dict["Kinetic Linkage"] = (f"{linkage}/100", "Elite" if linkage == 100 else "Optimal", "⛓️")
    elif any(x in sport for x in ["MARTIAL ARTS", "BOXING"]):
        snap = kpis.get("impact_snap", 0)
        m_dict["Impact Snap"] = (f"{snap:.1f}", "Elite" if snap > 15 else "Optimal", "🥊")
    else:
        duration = s['d1']['total'] / s['d1']['fps']
        m_dict["Movement Time"] = (f"{duration:.1f}s", "Consistent", "⏱️")
        
    return m_dict

# ============================================================================
# 6. ANALYTICS FUNCTIONS (Original)
# ============================================================================

def get_ai_metrics(raw_frames, fps):
    if not raw_frames: return None
    metrics = {
        "l_elbow": [], "r_elbow": [], "l_knee": [], "r_knee": [], "l_hip": [], "r_hip": [], 
        "wrist_speed": [], "hip_speed": [], "shoulder_speed": [], "ankle_speed": [],
        "trunk_lean": [], "trunk_lateral_lean": []
    }
    prev_w, prev_h, prev_s, prev_a = None, None, None, None
    for f in raw_frames:
        if not f:
            for k in metrics: metrics[k].append(None)
            continue
        metrics["l_elbow"].append(calculate_3d_angle(f[11], f[13], f[15]))
        metrics["r_elbow"].append(calculate_3d_angle(f[12], f[14], f[16]))
        metrics["l_knee"].append(calculate_3d_angle(f[23], f[25], f[27]))
        metrics["r_knee"].append(calculate_3d_angle(f[24], f[26], f[28]))
        metrics["l_hip"].append(calculate_3d_angle(f[11], f[23], f[25]))
        metrics["r_hip"].append(calculate_3d_angle(f[12], f[24], f[26]))
        
        curr_w = np.array([f[16]['x'], f[16]['y'], f[16]['z']])
        metrics["wrist_speed"].append(round(float(np.linalg.norm(curr_w - (prev_w if prev_w is not None else curr_w))*fps),4))
        prev_w = curr_w
        
        curr_h = np.array([(f[23]['x']+f[24]['x'])/2, (f[23]['y']+f[24]['y'])/2, (f[23]['z']+f[24]['z'])/2])
        metrics["hip_speed"].append(round(float(np.linalg.norm(curr_h - (prev_h if prev_h is not None else curr_h))*fps),4))
        prev_h = curr_h
        
        curr_s = np.array([(f[11]['x']+f[12]['x'])/2, (f[11]['y']+f[12]['y'])/2, (f[11]['z']+f[12]['z'])/2])
        metrics["shoulder_speed"].append(round(float(np.linalg.norm(curr_s - (prev_s if prev_s is not None else curr_s))*fps),4))
        prev_s = curr_s
        
        curr_a = np.array([f[28]['x'], f[28]['y'], f[28]['z']])
        metrics["ankle_speed"].append(round(float(np.linalg.norm(curr_a - (prev_a if prev_a is not None else curr_a))*fps),4))
        prev_a = curr_a
        
        mid_s, mid_h = get_midpoint(f[11], f[12]), get_midpoint(f[23], f[24])
        metrics["trunk_lean"].append(abs(calculate_lean(mid_s, mid_h, 'sagittal')))
        metrics["trunk_lateral_lean"].append(calculate_lean(mid_s, mid_h, 'coronal'))
    return metrics

def generate_sport_kpis(metrics, sport, raw_frames):
    kpis = {}
    def clean(lst): return [x if x is not None else 0 for x in lst]
    wrist_speed, hip_speed, shoulder_speed = clean(metrics["wrist_speed"]), clean(metrics["hip_speed"]), clean(metrics["shoulder_speed"])
    
    RACKET_BAT = ["TENNIS 🎾", "PADEL 🎾", "PICKLEBALL 🥒", "BADMINTON 🏸", "CRICKET 🏏", "GOLF ⛳", "BASEBALL ⚾"]
    if any(s in sport for s in RACKET_BAT) or "SOCCER" in sport:
        peak_hip, peak_shoulder, peak_wrist = np.argmax(hip_speed), np.argmax(shoulder_speed), np.argmax(wrist_speed)
        kpis["sequence_valid"] = peak_hip < peak_shoulder < peak_wrist
        timing_diff = peak_shoulder - peak_hip
        kpis["linkage_score"] = 100 if 2 <= timing_diff <= 10 else (70 if 0 < timing_diff < 2 else 40)
        
        if "GOLF" in sport:
            x_factors = []
            for f in raw_frames:
                if f:
                    s_vec = np.array([f[12]['x'] - f[11]['x'], f[12]['z'] - f[11]['z']])
                    h_vec = np.array([f[24]['x'] - f[23]['x'], f[24]['z'] - f[23]['z']])
                    cos_sim = np.dot(s_vec, h_vec) / (np.linalg.norm(s_vec) * np.linalg.norm(h_vec) + 1e-6)
                    x_factors.append(np.degrees(np.arccos(np.clip(cos_sim, -1.0, 1.0))))
                else: x_factors.append(0)
            kpis["max_x_factor"] = max(x_factors) if x_factors else 0
            
    elif "MARTIAL ARTS" in sport or "BOXING" in sport:
        peak_idx = np.argmax(wrist_speed)
        after_impact_idx = min(len(wrist_speed)-1, peak_idx + 2)
        kpis["impact_snap"] = abs(wrist_speed[peak_idx] - wrist_speed[after_impact_idx])
            
    elif "GYM" in sport:
        hip_y = [(f[23]['y'] + f[24]['y'])/2 for f in raw_frames if f]
        knee_y = [(f[25]['y'] + f[26]['y'])/2 for f in raw_frames if f]
        if hip_y and knee_y: kpis["depth_ratio"] = min(hip_y) / (max(knee_y) + 1e-6)
    elif "YOGA" in sport:
        com_x = [(f[23]['x'] + f[24]['x'])/2 for f in raw_frames if f]
        if com_x: kpis["stability"] = 1.0 - np.std(com_x)
    return kpis

def get_actionable_insights(kpis, sport):
    insights = []
    if "sequence_valid" in kpis:
        insights.append("✅ Perfect Kinetic Chain" if kpis["sequence_valid"] else "⚠️ Power Leak: Lead with the hips.")
    if sport == "GOLF ⛳" and "max_x_factor" in kpis:
        insights.append("✅ Elite Rotation" if kpis["max_x_factor"] > 30 else "💡 Increase X-Factor separation.")
    if "GYM" in sport:
        if kpis.get("depth_ratio", 0) > 0.9: insights.append("💡 Depth: Lower hips break parallel.")
    if "YOGA" in sport and "stability" in kpis:
        insights.append("✅ Zen Stability" if kpis["stability"] > 0.98 else "💡 Core Focus: Minimize lateral wobble.")
    if not insights: insights = ["Focus on consistent tempo.", "Keep your core engaged.", "Maintain visual focus."]
    return insights[:3]

# ============================================================================
# 7. CHARTING & RENDERING (Original)
# ============================================================================

def plot_power_curve(metrics):
    def clean(lst): return [x if x is not None else 0 for x in lst]
    fig = go.Figure()
    fig.add_trace(go.Scatter(y=clean(metrics["wrist_speed"]), fill='tozeroy', name="Wrist Velocity", line=dict(color='#ccff00')))
    fig.update_layout(template="plotly_dark", title="Power Curve", margin=dict(l=20, r=20, t=40, b=20), height=300)
    return fig

def plot_radar_chart(metrics):
    def clean_max(lst): 
        vals = [x for x in lst if x is not None]
        return max(vals) if vals else 0
    categories = ['L Elbow', 'R Elbow', 'L Knee', 'R Knee', 'L Hip', 'R Hip']
    user_vals = [clean_max(metrics[c.lower().replace(' ', '_')]) for c in categories]
    pro_vals = [160, 160, 140, 140, 120, 120]
    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(r=user_vals, theta=categories, fill='toself', name='User'))
    fig.add_trace(go.Scatterpolar(r=pro_vals, theta=categories, fill='toself', name='Pro Benchmark', line=dict(color='#ccff00')))
    fig.update_layout(polar=dict(radialaxis=dict(visible=True, range=[0, 180])), template="plotly_dark", margin=dict(l=40, r=40, t=40, b=40), height=300)
    return fig

def plot_kinetic_chain(metrics):
    def clean(lst): return [x if x is not None else 0 for x in lst]
    fig = go.Figure()
    fig.add_trace(go.Scatter(y=clean(metrics["hip_speed"]), name="Hip", stackgroup='one', line=dict(color='#1e293b')))
    fig.add_trace(go.Scatter(y=clean(metrics["shoulder_speed"]), name="Torso", stackgroup='one', line=dict(color='#38bdf8')))
    fig.add_trace(go.Scatter(y=clean(metrics["wrist_speed"]), name="Arm", stackgroup='one', line=dict(color='#ccff00')))
    fig.update_layout(template="plotly_dark", title="Kinetic Chain", margin=dict(l=20, r=20, t=40, b=20), height=300)
    return fig

def download_model():
    p = 'pose_landmarker_heavy.task'
    if not os.path.exists(p): urllib.request.urlretrieve("https://storage.googleapis.com/mediapipe-models/pose_landmarker/pose_landmarker_heavy/float16/1/pose_landmarker_heavy.task", p)
    return p

def analyze_vid(path, model):
    det = vision.PoseLandmarker.create_from_options(vision.PoseLandmarkerOptions(base_options=python.BaseOptions(model_asset_path=model), running_mode=vision.RunningMode.VIDEO))
    cap = cv2.VideoCapture(path)
    fps, history, raw, impact_f, peak_v, prev_w = cap.get(cv2.CAP_PROP_FPS), [], [], 0, 0, None
    while cap.isOpened():
        ret, frame = cap.read()
        if not ret: break
        ts = int((len(history) * 1000) / (fps if fps > 0 else 30))
        res = det.detect_for_video(mp.Image(image_format=mp.ImageFormat.SRGB, data=cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)), ts)
        lms = res.pose_landmarks[0] if res.pose_landmarks else None
        history.append(lms)
        raw.append([{"id": j, "x": l.x, "y": l.y, "z": l.z} for j, l in enumerate(lms)] if lms else None)
        if lms and res.pose_world_landmarks:
            w = res.pose_world_landmarks[0][15]
            if prev_w:
                v = np.linalg.norm(np.array([w.x, w.y, w.z]) - np.array([prev_w.x, prev_w.y, prev_w.z]))
                if v > peak_v: peak_v, impact_f = v, len(history)-1
            prev_w = w
    cap.release(); return {"history": history, "raw": raw, "fps": fps, "total": len(history), "impact": impact_f}

def draw_neon_skeleton(img, lms, alpha=0.8):
    if not lms: return
    overlay = img.copy()
    # High-detail skeletal structure including hands and feet
    FULL_SKELETON = [
        (11,12), (11,13), (13,15), (12,14), (14,16), # Shoulders and arms
        (15,17), (15,19), (15,21), (17,19), # Left hand
        (16,18), (16,20), (16,22), (18,20), # Right hand
        (11,23), (12,24), (23,24), # Torso
        (23,25), (25,27), (24,26), (26,28), # Legs
        (27,29), (27,31), (29,31), # Left foot
        (28,30), (28,32), (30,32)  # Right foot
    ]
    # Draw connections with a dark outer glow for contrast
    for s, e in FULL_SKELETON:
        p1 = (int(lms[s].x*img.shape[1]), int(lms[s].y*img.shape[0]))
        p2 = (int(lms[e].x*img.shape[1]), int(lms[e].y*img.shape[0]))
        cv2.line(overlay, p1, p2, (0, 0, 0), 3, cv2.LINE_AA) # Outer glow
        cv2.line(overlay, p1, p2, (255, 255, 255), 1, cv2.LINE_AA) # White core
    
    # Draw joints (Neon Green with White core)
    for i in range(11, 33):
        pt = (int(lms[i].x*img.shape[1]), int(lms[i].y*img.shape[0]))
        cv2.circle(overlay, pt, 4, (0, 0, 0), -1, cv2.LINE_AA) # Border
        cv2.circle(overlay, pt, 3, (0, 255, 0), -1, cv2.LINE_AA) # Green
        cv2.circle(overlay, pt, 1, (255, 255, 255), -1, cv2.LINE_AA) # White core
        
    cv2.addWeighted(overlay, alpha, img, 1 - alpha, 0, img)

def render_pro_stereo(p1, p2, h1, h2, f1, f2, fps):
    cap1 = cv2.VideoCapture(p1)
    target_h = 720
    w1 = int(cap1.get(3)*(target_h/cap1.get(4)))
    if p2:
        cap2 = cv2.VideoCapture(p2)
        w2 = int(cap2.get(3)*(target_h/cap2.get(4)))
        combined_w, off = w1 + w2, f1 - f2
    else: combined_w, off = w1, 0
    raw_p, final_p = os.path.join(tempfile.gettempdir(), f"r_{int(time.time())}.mp4"), os.path.join(tempfile.gettempdir(), f"p_{int(time.time())}.mp4")
    out = cv2.VideoWriter(raw_p, cv2.VideoWriter_fourcc(*'mp4v'), fps, (combined_w, target_h))
    for i in range(len(h1)):
        ret1, f1_img = cap1.read()
        if not ret1: break
        if h1[i]: draw_neon_skeleton(f1_img, h1[i], alpha=0.8) # High visibility
        frame_to_write = cv2.resize(f1_img, (w1, target_h))
        if p2:
            idx2 = i - off
            if 0 <= idx2 < len(h2):
                cap2.set(1, idx2); _, f2_img = cap2.read()
                if h2[idx2]: draw_neon_skeleton(f2_img, h2[idx2], alpha=0.8) # High visibility
                f2_img = cv2.resize(f2_img, (w2, target_h))
            else: f2_img = np.zeros((target_h, w2, 3), dtype=np.uint8)
            frame_to_write = np.hstack((frame_to_write, f2_img))
        out.write(frame_to_write)
    cap1.release()
    if p2: cap2.release()
    out.release()
    subprocess.run(f'ffmpeg -y -i "{raw_p}" -c:v libx264 -pix_fmt yuv420p -preset ultrafast "{final_p}"', shell=True)
    return final_p

# ============================================================================
# 8. SPORT CONFIG & MAIN APP (Merged logic)
# ============================================================================

SPORT_CONFIG = {
    "TENNIS 🎾": ["General Rally", "First Serve", "Second Serve", "Forehand Flat", "Forehand Topspin", "Forehand Slice", "Backhand One-Handed", "Backhand Two-Handed", "Backhand Slice", "Forehand Volley", "Backhand Volley", "Overhead Smash", "Drop Shot", "Lob", "Return of Serve"],
    "PADEL 🎾": ["General Rally", "Serve", "Forehand Groundstroke", "Backhand Groundstroke", "Bandeja", "Vibora", "Flat Smash", "Smash 'por 3'", "Smash 'por 4'", "Forehand Volley", "Backhand Volley", "Bajada de Pared", "Chiquita", "Globo (Lob)", "Contrapared"],
    "PICKLEBALL 🥒": ["General Rally", "Serve (Volley)", "Serve (Drop)", "Dink (Straight)", "Dink (Cross-court)", "Third Shot Drop", "Third Shot Drive", "Speed Up", "Kitchen Volley", "Punch Volley", "Overhead Slam", "Reset Shot", "Backhand Flick", "Lob"],
    "GOLF ⛳": ["Driver Tee Shot", "Fairway Wood", "Long Iron", "Short Iron", "Pitch Shot", "Chipping", "Sand Bunker Shot", "Putter Stroke", "Full Backswing", "Downswing Transition", "Follow Through"],
    "BADMINTON 🏸": ["General Rally", "High Serve", "Low Serve", "Flick Serve", "Forehand Smash", "Backhand Smash", "Jump Smash", "Clear (Lob)", "Drop Shot", "Net Kill", "Net Lift", "Drive Shot", "Around-the-head Shot"],
    "CRICKET 🏏": ["Forward Defense", "Cover Drive", "Pull Shot", "Hook Shot", "Cut Shot", "Sweep Shot", "Fast Bowling Action", "Spin Bowling Action", "Wicket-keeping Stance", "Power Hitting (Slog)", "High Catching"],
    "SOCCER ⚽": ["Instep Drive (Power)", "Side-foot Pass", "Curled Shot", "Long Ball/Switch", "Heading (Standing)", "Heading (Jumping)", "Goalkeeper Dive", "Goalkeeper Goal Kick", "Throw-in", "Penalty Kick", "Volley", "First Touch Control"],
    "BASKETBALL 🏀": ["Jump Shot", "Three-Pointer", "Free Throw", "Layup (Right/Left)", "Driving Dunk", "Chest Pass", "Bounce Pass", "Overhead Pass", "Defensive Slide", "Post-up Turnaround", "Rebounding Box-out"],
    "BASEBALL ⚾": ["Pitching (Wind-up)", "Pitching (Stretch)", "Power Swing", "Bunt", "Catcher Throw-down", "Infield Scoop", "Sliding"],
    "AMERICAN FOOTBALL 🏈": ["QB Drop-back Pass", "QB Shotgun Pass", "Field Goal Kick", "Punting", "WR Route Cut", "Lineman Drive Block"],
    "ICE HOCKEY 🏒": ["Slap Shot", "Wrist Shot", "Snap Shot", "Backhand Shot", "Skating Stride (Start)", "Skating Crossover"],
    "TABLE TENNIS 🏓": ["Forehand Loop", "Backhand Push", "Pendulum Serve", "Forehand Smash", "Backhand Flick", "Chop"],
    "MARTIAL ARTS 🥋": ["Jab-Cross Combo", "Roundhouse Kick", "Front Kick", "Lead Hook", "Double Leg Takedown", "Sprawl", "Block/Parry"],
    "BOXING/MMA 🥊": ["Jab", "Cross", "Lead Hook", "Rear Hook", "Uppercut", "Lead Roundhouse Kick", "Rear Roundhouse Kick", "Front Kick (Teep)", "Shoulder Roll", "Slip/Bob and Weave", "Double Leg Takedown", "Sprawl"],
    "GYM 🏋️": ["Back Squat", "Front Squat", "Deadlift (Conventional)", "Deadlift (Sumo)", "Bench Press", "Overhead Press", "Barbell Row", "Pull-up", "Clean and Jerk", "Snatch", "Kettlebell Swing", "Lunge"],
    "YOGA 🧘": ["Downward Dog", "Warrior I", "Warrior II", "Warrior III", "Tree Pose", "Triangle Pose", "Crow Pose", "Plank/Chaturanga", "Cobra/Upward Dog", "Half Moon Pose", "Bridge Pose", "Wheel Pose"],
    "ATHLETICS/RUNNING 🏃": ["Sprint Start (Blocks)", "Max Velocity Phase", "Distance Running Gate", "Long Jump Takeoff", "High Jump Fosbury Flop", "Hurdle Clearance", "Shot Put Glide", "Shot Put Rotational", "Javelin Throw"]
}

# App Header
st.markdown("<h1>Vector Victor AI</h1>", unsafe_allow_html=True)
st.markdown("<p class='hero-sub'>Deep form Vector based Bio mechanics AI engine</p>", unsafe_allow_html=True)

# Setup Gemini
if "GEMINI_API_KEY" in st.secrets:
    genai.configure(api_key=st.secrets["GEMINI_API_KEY"])

import plotly.graph_objects as go
import plotly.express as px

st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Bitcount+Prop+Single&display=swap');
    :root {
        --neon-green: #ccff00;
        --glass-bg: rgba(255, 255, 255, 0.05);
        --glass-border: rgba(255, 255, 255, 0.1);
    }
    
    /* Global Font Override */
    .stApp, .stApp * { 
        font-family: 'Bitcount Prop Single', sans-serif !important; 
    }
    
    /* Global Uppercase for Interactive Elements */
    button, label, input, select, textarea, .hero-sub, h4, [data-testid="stWidgetLabel"] p { 
        text-transform: uppercase !important;
    }

    /* Standardise File Uploader - Keep Font, Fix Overlap by disabling uppercase */
     /* Target the 'Browse files' button text */
    [data-testid="stFileUploader"] section button div {
        font-size: 0 !important; /* Hides the original 'Browse files' text */
    }
    
    /* Inject your custom name */
    [data-testid="stFileUploader"] section button div::before {
        content: "_                _"; /* Replace this with " " if you want it empty */
        font-size: 14px !important;
        font-family: 'Bitcount Prop Single', sans-serif !important;
        text-transform: uppercase;
        visibility: visible;
    }

/* Remove the 'Limit 200MB' and original filename helper text if you want a cleaner look */
[data-testid="stFileUploader"] section > div {
    display: none !important;
}
    
    .stApp { background: radial-gradient(circle at top right, #0f172a, #020617); color: #f8fafc; }
    .main { padding: 1rem 0.5rem; }
    
    /* Button Styling */
    button { 
        min-height: 48px !important;
        font-size: 14px !important;
    }
    .stButton > button { width: 100% !important; }

    /* Tabs Styling */
    .stTabs [data-baseweb="tab-list"] { gap: 0.5rem; display: flex; flex-wrap: wrap; justify-content: center; }
    .stTabs [data-baseweb="tab"] {
        flex: 1; min-width: 80px; text-align: center; height: 50px; 
        background: rgba(255, 255, 255, 0.05) !important; border-radius: 10px !important; 
        border: 1px solid rgba(255,255,255,0.1) !important;
    }
    .stTabs [aria-selected="true"] { background: rgba(204, 255, 0, 0.1) !important; border: 1px solid var(--neon-green) !important; }
    .stTabs [aria-selected="true"] p { color: var(--neon-green) !important; font-weight: 700 !important; }

    .stSlider { padding: 1rem 0; }
    [data-testid="column"] { padding: 0 0.25rem; }
    div[data-testid="stSlider"] label p { color: var(--neon-green) !important; font-weight: 900 !important; font-size: 1rem !important; letter-spacing: 1px; }
    div[data-testid="stThumbValue"] { color: #000 !important; background-color: var(--neon-green) !important; font-weight: 900 !important; }
    div[data-baseweb="slider"] > div { background: var(--neon-green) !important; }
    
    .glass-card { background: var(--glass-bg); backdrop-filter: blur(12px); border: 1px solid var(--glass-border); border-radius: 24px; padding: 1.5rem; margin-bottom: 2rem; }
    .bento-card { background: rgba(255, 255, 255, 0.03); border: 1px solid var(--glass-border); border-radius: 20px; padding: 20px; text-align: center; margin-bottom: 10px; }

    /* Heading Styling */
    h1 { font-size: clamp(1.5rem, 6vw, 3rem) !important; font-weight: normal !important; background: linear-gradient(to right, #00f2fe, #ff00ff); -webkit-background-clip: text; -webkit-text-fill-color: transparent; text-align: center; margin-bottom: 0px !important; text-transform: uppercase !important; }
    h4 { font-size: 0.85rem !important; font-weight: 700 !important; color: var(--neon-green); margin-top: 2.5rem !important; margin-bottom: 0.5rem !important; letter-spacing: 2px; }
    .hero-sub { text-align: center; color: #94a3b8; font-size: 0.39375rem; letter-spacing: 4px; margin-bottom: 2rem; }

    /* Specific Button Highlighting */
    div.stButton > button:has(div:contains("GENERATE AI COACHING REPORT")) { background: linear-gradient(135deg, #ffd700 0%, #daa520 50%, #b8860b 100%) !important; color: #000 !important; font-weight: 900 !important; }
    </style>
    """, unsafe_allow_html=True)

# Setup Gemini
if "GEMINI_API_KEY" in st.secrets:
    genai.configure(api_key=st.secrets["GEMINI_API_KEY"])

import plotly.graph_objects as go
import plotly.express as px

class NpEncoder(json.JSONEncoder):
    def default(self, obj):
        if isinstance(obj, np.integer): return int(obj)
        if isinstance(obj, np.floating): return float(obj)
        if isinstance(obj, np.bool_): return bool(obj)
        if isinstance(obj, np.ndarray): return obj.tolist()
        return super(NpEncoder, self).default(obj)

# Tabs Navigation
tab1, tab2, tab3 = st.tabs(["UPLOAD", "ANALYZE", "RESULTS"])

# Tab 1: Upload
with tab1:
    st.markdown("#### UPLOAD VIDEOS", unsafe_allow_html=True)
    selected_sport = st.selectbox("SELECT SPORT", list(SPORT_CONFIG.keys()), key="sport_sel")
    selected_action = st.selectbox("SELECT ACTION", SPORT_CONFIG[selected_sport], key="action_sel")
    
    st.markdown("#### PRIMARY VIEW", unsafe_allow_html=True)
    u1 = st.file_uploader("UPLOAD MAIN ANGLE", type=["mp4","mov"], key="u1_upload", label_visibility="collapsed")
    if u1: display_file_info(u1)
    
    is_stereo = st.toggle("STEREOGRAPHIC MODE (DUAL VIEW)", value=False, key="st_toggle")
    u2 = None
    if is_stereo:
        st.markdown("#### SECONDARY VIEW", unsafe_allow_html=True)
        u2 = st.file_uploader("UPLOAD SECOND ANGLE", type=["mp4","mov"], key="u2_upload", label_visibility="collapsed")
        if u2: display_file_info(u2)

    if u1:
        if st.button("PROCEED TO ANALYSIS", type="primary", width="stretch"):
            st.session_state["u1"] = u1
            st.session_state["u2"] = u2
            st.session_state["sport"] = selected_sport
            st.session_state["action"] = selected_action

            # Clear previous auto-detection data
            if "manual_actions" in st.session_state: del st.session_state["manual_actions"]
            if "tele_opt_preview" in st.session_state: del st.session_state["tele_opt_preview"]
            if "tele_opt" in st.session_state: del st.session_state["tele_opt"]
            if "report_text" in st.session_state: del st.session_state["report_text"]
            if "final_video" in st.session_state: del st.session_state["final_video"]

            model_task = download_model()
            t1_p = os.path.join(tempfile.gettempdir(), f"l_raw.mp4")
            with open(t1_p, "wb") as f: f.write(u1.getbuffer())
            
            with st.status(f"ANALYZING {selected_sport.upper()}...") as status:
                d1 = analyze_vid(t1_p, model_task)
                d2, t2_p = None, None
                if is_stereo and u2:
                    t2_p = os.path.join(tempfile.gettempdir(), f"s_raw.mp4")
                    with open(t2_p, "wb") as f: f.write(u2.getbuffer())
                    d2 = analyze_vid(t2_p, model_task)
                
                st.session_state["data_current"] = {"d1": d1, "d2": d2, "p1": t1_p, "p2": t2_p}
                status.update(label="INITIAL ANALYSIS COMPLETE!", state="complete")
            
            # INFO FOR USER
            st.success("✅ INITIAL ANALYSIS COMPLETE! PLEASE PROCEED TO THE 'ANALYZE' TAB ABOVE.")
    else:
        st.info("📌 UPLOAD A VIDEO TO START.")

# Tab 2: Analyze
with tab2:
    if "data_current" not in st.session_state:
        st.warning("⚠️ UPLOAD AND PROCESS A VIDEO FIRST.")
    else:
        s = st.session_state["data_current"]
        sport, action = st.session_state["sport"], st.session_state["action"]
        st.markdown(f"#### SYNCHRONIZE: {sport.upper()}", unsafe_allow_html=True)
        
        st.markdown("#### SOURCE 1 SYNC")
        sl1 = st.slider("ALIGN ON IMPACT FRAME", 0, s['d1']['total']-1, s['d1']['impact'], key="sl1_sync", label_visibility="collapsed")
        st.caption(f"IMPACT FRAME: {sl1}")
        
        sl2 = 0
        if s['p2']:
            st.markdown("#### SOURCE 2 SYNC")
            sl2 = st.slider("ALIGN SECONDARY VIEW", 0, s['d2']['total']-1, s['d2']['impact'], key="sl2_sync", label_visibility="collapsed")
            st.caption(f"SECONDARY SYNC: {sl2}")
            
        detected_side = detect_handedness(s['d1']['raw'])
        handedness_choice = st.radio("ATHLETE HANDEDNESS", ["Auto-detect (Detected: " + detected_side.upper() + ")", "Left Handed", "Right Handed"], index=0, key="hand_sel")
        
        # Preview
        cap1 = cv2.VideoCapture(s['p1']); cap1.set(cv2.CAP_PROP_POS_FRAMES, sl1)
        ret1, i1 = cap1.read(); cap1.release()
        if ret1 and i1 is not None:
            i1 = cv2.cvtColor(i1, cv2.COLOR_BGR2RGB)
            if s['p2']:
                cap2 = cv2.VideoCapture(s['p2']); cap2.set(cv2.CAP_PROP_POS_FRAMES, sl2)
                ret2, i2 = cap2.read(); cap2.release()
                if ret2 and i2 is not None:
                    i2 = cv2.cvtColor(i2, cv2.COLOR_BGR2RGB)
                    h_t = 400
                    w1_n, w2_n = int(i1.shape[1] * (h_t / i1.shape[0])), int(i2.shape[1] * (h_t / i2.shape[0]))
                    st.image(np.hstack((cv2.resize(i1, (w1_n, h_t)), cv2.resize(i2, (w2_n, h_t)))), width="stretch")
            else:
                st.image(i1, width="stretch")
        
        # --- AUTO-DETECTED ACTIONS (For General Rally) ---
        sport_clean = "".join([c for c in sport if ord(c) < 128]).strip().upper()
        action_upper = action.upper()
        
        if sport_clean == "TENNIS" and action_upper == "GENERAL RALLY":
            with st.expander("🤖 AUTO-DETECTION SETTINGS (Beta)", expanded=True):
                use_detection = st.toggle("Enable multi-motion detection", value=True, key="use_multi_detect")
                
                if use_detection:
                    # Need metrics for auto-detection
                    if "tele_opt_preview" not in st.session_state:
                        with st.spinner("SCANNING FOR MOTIONS..."):
                            raw_interp = interpolate_landmarks(s['d1']['raw'])
                            # Build a preview telemetry object to get metrics
                            st.session_state["tele_opt_preview"] = build_pro_telemetry(raw_interp, sport, action, sl1, s['d1']['fps'], "lead")
                    
                    # Initialize session state for actions if not present
                    if "manual_actions" not in st.session_state:
                        st.session_state.manual_actions = auto_detect_actions(
                            st.session_state["tele_opt_preview"]["metrics"], sport_clean, fps=s['d1']['fps']
                        )

                    # Option to add a new manual action
                    if st.button("➕ Add Manual Motion"):
                        st.session_state.manual_actions.append({"action": "FOREHAND DRIVE", "frame": sl1})
                        st.rerun()

                    if st.session_state.manual_actions:
                        # Create a list to track indices to remove
                        to_remove = []
                        
                        for idx, act in enumerate(st.session_state.manual_actions):
                            c1, c2, c3 = st.columns([2, 2, 1])
                            with c1:
                                # Allow editing the action type
                                available_actions = list(ACTION_SIGNATURES["TENNIS"].keys())
                                current_idx = 0
                                if act['action'] in available_actions:
                                    current_idx = available_actions.index(act['action'])
                                
                                act['action'] = st.selectbox(f"Type", available_actions, index=current_idx, key=f"type_{idx}", label_visibility="collapsed")
                            with c2:
                                act['frame'] = st.number_input(f"Frame", value=int(act['frame']), step=1, key=f"act_{idx}", label_visibility="collapsed")
                            with c3:
                                if st.button("🗑️", key=f"del_{idx}"):
                                    to_remove.append(idx)
                        
                        # Perform removals if any
                        if to_remove:
                            for index in sorted(to_remove, reverse=True):
                                st.session_state.manual_actions.pop(index)
                            st.rerun()
                    else:
                        st.info("No specific shots auto-detected. Use 'Add Manual Motion' to define key frames.")
                else:
                    st.info("Multi-motion detection disabled. Analysis will focus on the single impact frame selected above.")

        if st.button("🚀 START FINAL BIOMECHANICAL RENDER", type="primary", width="stretch"):
            with st.spinner("PROCESSING VECTORS..."):
                final_v = render_pro_stereo(s['p1'], s['p2'], s['d1']['history'], (s['d2']['history'] if s['d2'] else []), sl1, sl2, s['d1']['fps'])
                st.session_state["final_video"] = final_v
                
                h_val = None
                if "Left" in st.session_state.hand_sel: h_val = "left"
                elif "Right" in st.session_state.hand_sel: h_val = "right"
                
                raw_interp = interpolate_landmarks(s['d1']['raw'])
                tele_opt = build_pro_telemetry(raw_interp, sport, action, sl1, s['d1']['fps'], "dual" if s['p2'] else "lead", handedness_override=h_val)
                
                # --- INJECT DETECTED ACTIONS (If enabled) ---
                if st.session_state.get("use_multi_detect", False) and "manual_actions" in st.session_state:
                    tele_opt["metadata"]["detected_actions"] = st.session_state.manual_actions
                
                st.session_state["tele_opt"] = tele_opt
                #st.session_state["brief"] = generate_brief(tele_opt)
                st.session_state["brief"] = generate_brief.generate_brief(
					tele_opt,
					sport=sport,
					action=action
				)
                st.session_state["sl1_val"] = sl1 # For efficiency calc
            
            # INFO FOR USER
            st.success("✅ BIOMECHANICAL RENDER COMPLETE! PLEASE PROCEED TO THE 'RESULTS' TAB ABOVE.")

# Tab 3: Results
with tab3:
    if "final_video" not in st.session_state:
        st.warning("⚠️ COMPLETE THE SYNCHRONIZATION AND RENDER FIRST.")
    else:
        s = st.session_state["data_current"]
        sport, action = st.session_state["sport"], st.session_state["action"]
        st.markdown("#### PERFORMANCE ANALYSIS", unsafe_allow_html=True)
        st.video(st.session_state["final_video"])
        
        st.markdown("#### PRO ANALYTICS DASHBOARD")
        metrics = get_ai_metrics(s['d1']['raw'], s['d1']['fps'])
        if metrics:
            kpis = generate_sport_kpis(metrics, sport, s['d1']['raw'])
            insights = get_actionable_insights(kpis, sport)
            m_grid = get_sport_metrics(sport, metrics, st.session_state["sl1_val"], kpis, s)
            draw_mobile_metric_grid(m_grid)
            
            st.markdown("---")
            st.markdown("#### AI COACHING INSIGHTS")
            for insight in insights: st.success(insight.upper())
            
            st.markdown("---")
            st.markdown("#### DETAILED ANALYTICS")
            chart_view = st.radio("CHOOSE VIEW", ["POWER CURVE", "RADAR CHART", "KINETIC CHAIN"], horizontal=True, label_visibility="collapsed")
            if chart_view == "POWER CURVE": st.plotly_chart(plot_power_curve(metrics), width="stretch")
            elif chart_view == "RADAR CHART": st.plotly_chart(plot_radar_chart(metrics), width="stretch")
            else: st.plotly_chart(plot_kinetic_chain(metrics), width="stretch")
            
            # --- DOWNLOADS & EXPORTS ---
            st.markdown("#### DOWNLOADS & EXPORT")
            
            # ZIP Download (Always available after render)
            z_buf = io.BytesIO()
            with zipfile.ZipFile(z_buf, "w") as zf:
                zf.write(st.session_state["final_video"], "analysis.mp4")
                zf.writestr("AI_BRIEF.txt", st.session_state["brief"])
            
            cz1, cz2 = st.columns(2)
            with cz1:
                st.download_button("📥 DOWNLOAD ZIP (VIDEO + DATA)", z_buf.getvalue(), f"{sport}_DATA.zip", width="stretch")
            with cz2:
                json_data = json.dumps(st.session_state["tele_opt"], indent=2, cls=NpEncoder)
                st.download_button("💾 RAW JSON", json_data, f"{sport}_TELEMETRY.json", "application/json", width="stretch")
            
            st.markdown("---")
            if st.button("🤖 GENERATE AI COACHING REPORT", type="primary", width="stretch"):
                with st.status("AI IS ANALYZING...") as status:
                    #report_text = generate_pro_report(st.session_state["brief"])
                    report_text = generate_pro_report(
                    st.session_state["brief"],
                    sport=st.session_state["sport"],
                    action=st.session_state["action"]
                    )
                    st.session_state["report_text"] = report_text
                    status.update(label="REPORT COMPLETE!", state="complete")
            
            if "report_text" in st.session_state:
                st.markdown(st.session_state["report_text"])
                st.markdown("#### EXPORT DOCUMENTS")
                c1, c2 = st.columns(2)
                
                # Get hand dominance
                hand = st.session_state["tele_opt"]["metadata"].get("dominant_side", "unknown")
                
                with c1:
                    docx_f = create_docx_report(st.session_state["report_text"], sport, action, hand)
                    st.download_button("📄 WORD DOC", docx_f, f"{sport}_ANALYSIS.docx", width="stretch")
                with c2:
                    pdf_f = create_pdf_report(st.session_state["report_text"], sport, action, hand)
                    st.download_button("📜 PDF REPORT", pdf_f, f"{sport}_ANALYSIS.pdf", width="stretch")

        # ====================================================================
        # IMPACT REPLAY RENDERER
        # ====================================================================
        if "tele_opt" in st.session_state:
            st.markdown("---")
            st.markdown("### 🎥 IMPACT REPLAY SETUP")
            st.info("SCROLL THE SLIDER TO SELECT THE EXACT IMPACT FRAME ON THE FINAL RENDERED VIDEO.")
            
            fps = st.session_state["tele_opt"]["metadata"].get("fps", 30)
            input_vid = st.session_state["final_video"]
            
            # Use OpenCV to get total frame count for the slider
            cap = cv2.VideoCapture(input_vid)
            total_frames = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
            cap.release()
            
            # Default to the impact frame aligned in the Analyze tab
            default_impact = st.session_state.get("sl1_val", 0)
            impact_f = st.slider("SELECT IMPACT FRAME", 0, total_frames - 1, int(default_impact), key="impact_replay_slider")
            
            # Preview the selected frame
            cap = cv2.VideoCapture(input_vid)
            cap.set(cv2.CAP_PROP_POS_FRAMES, impact_f)
            ret, frame = cap.read()
            cap.release()
            if ret:
                st.image(cv2.cvtColor(frame, cv2.COLOR_BGR2RGB), caption=f"PREVIEW: FRAME {impact_f}", width="stretch")
                
            slow_mo_output = os.path.join(tempfile.gettempdir(), f"replay_{int(time.time())}.mp4")

            if st.button("🎬 RENDER 0.2X SLOW-MO REPLAY", type="secondary", width="stretch"):
                with st.spinner("RENDERING SMOOTH SLOW-MO (MINTERPOLATE)... THIS MAY TAKE A MOMENT."):
                    try:
                        result = apply_advanced_slow_mo(input_vid, slow_mo_output, impact_f, fps)
                        st.session_state["slow_mo_video"] = result
                    except Exception as e:
                        st.error(f"REPLAY RENDERING ERROR: {e}")

            if "slow_mo_video" in st.session_state:
                st.markdown("#### ✅ GENERATED REPLAY")
                st.video(st.session_state["slow_mo_video"])
                with open(st.session_state["slow_mo_video"], "rb") as f:
                    st.download_button("📥 DOWNLOAD REPLAY", f, f"{sport}_REPLAY.mp4", width="stretch")

        st.markdown("---")
        if st.button("↺ ANALYZE ANOTHER VIDEO", width="stretch"):
            for key in list(st.session_state.keys()): del st.session_state[key]
            st.rerun()
